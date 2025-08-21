import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

def parse_junit_xml(xml_path):
    tree = ET.parse(xml_path)
    root = tree.getroot()
    results = []
    suite_name = root.attrib.get("name", "TestSuite")
    for testcase in root.findall("testcase"):
        name = testcase.attrib.get("name", "")
        classname = testcase.attrib.get("classname", "")
        time = testcase.attrib.get("time", "0")
        status = "passed"
        message = ""
        details = ""
        if testcase.find("failure") is not None:
            status = "failed"
            failure = testcase.find("failure")
            message = failure.attrib.get("message", "")
            details = failure.text or ""
        elif testcase.find("error") is not None:
            status = "error"
            error = testcase.find("error")
            message = error.attrib.get("message", "")
            details = error.text or ""
        elif testcase.find("skipped") is not None:
            status = "skipped"
        results.append({
            "name": name,
            "classname": classname,
            "time": time,
            "status": status,
            "message": message,
            "details": details
        })
    summary = {
        "suite_name": suite_name,
        "total": int(root.attrib.get("tests", len(results))),
        "failures": int(root.attrib.get("failures", 0)),
        "errors": int(root.attrib.get("errors", 0)),
        "skipped": int(root.attrib.get("skipped", 0)),
        "time": root.attrib.get("time", "0"),
        "results": results
    }
    return summary

def create_docx_report(summary, output_path):
    doc = Document()
    doc.add_heading(f"JUnit Test Report: {summary['suite_name']}", 0)

    # Summary table
    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Total"
    hdr_cells[1].text = "Failures"
    hdr_cells[2].text = "Errors"
    hdr_cells[3].text = "Skipped"
    hdr_cells[4].text = "Time (s)"
    row = table.add_row().cells
    row[0].text = str(summary['total'])
    row[1].text = str(summary['failures'])
    row[2].text = str(summary['errors'])
    row[3].text = str(summary['skipped'])
    row[4].text = summary['time']

    doc.add_paragraph("")

    # Test case details
    doc.add_heading("Test Cases", level=1)
    for result in summary['results']:
        p = doc.add_paragraph()
        status = result['status']
        run_time = result['time']
        p.add_run(f"{result['name']} ({result['classname']}) - {status.upper()} [Time: {run_time}s]").bold = True
        if status == "failed":
            run = p.add_run(f"\n  Failure: {result['message']}")
            run.font.color.rgb = RGBColor(255, 0, 0)      # Red
            if result['details']:
                p.add_run(f"\n  Details: {result['details']}")
        elif status == "error":
            run = p.add_run(f"\n  Error: {result['message']}")
            run.font.color.rgb = RGBColor(255, 165, 0)    # Orange
            if result['details']:
                p.add_run(f"\n  Details: {result['details']}")
        elif status == "skipped":
            run = p.add_run("\n  Skipped")
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    doc.save(output_path)

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Parse JUnit XML and generate DOCX report")
    parser.add_argument("xml_file", help="Path to the JUnit XML file")
    parser.add_argument("output_file", help="Path to the output DOCX file")
    args = parser.parse_args()
    summary = parse_junit_xml(args.xml_file)
    create_docx_report(summary, args.output_file)
    print(f"Report saved to {args.output_file}")