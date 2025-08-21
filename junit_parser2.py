#!/usr/bin/env python3
"""
JUnit XML to Word Report Generator

This script parses JUnit XML test results and generates a professional-looking
Word document report using python-docx.

Requirements:
    pip install python-docx lxml

Usage:
    python junit_to_word_report.py input.xml output.docx
"""

import argparse
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
import sys
from dataclasses import dataclass
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import RGBColor
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


@dataclass
class TestCase:
    """Represents a single test case from JUnit XML."""
    name: str
    classname: str
    time: float
    status: str  # 'passed', 'failed', 'skipped', 'error'
    failure_message: Optional[str] = None
    failure_details: Optional[str] = None
    skipped_message: Optional[str] = None


@dataclass
class TestSuite:
    """Represents a test suite from JUnit XML."""
    name: str
    tests: int
    failures: int
    errors: int
    skipped: int
    time: float
    test_cases: List[TestCase]


class JUnitXMLParser:
    """Parser for JUnit XML files."""
    
    def parse(self, xml_file: Path) -> List[TestSuite]:
        """Parse JUnit XML file and return list of test suites."""
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()
        except ET.ParseError as e:
            raise ValueError(f"Invalid XML file: {e}")
        
        test_suites = []
        
        # Handle both single testsuite and testsuites root elements
        if root.tag == 'testsuite':
            test_suites.append(self._parse_testsuite(root))
        elif root.tag == 'testsuites':
            for testsuite in root.findall('testsuite'):
                test_suites.append(self._parse_testsuite(testsuite))
        else:
            raise ValueError(f"Unexpected root element: {root.tag}")
        
        return test_suites
    
    def _parse_testsuite(self, testsuite_elem) -> TestSuite:
        """Parse a single testsuite element."""
        name = testsuite_elem.get('name', 'Unknown')
        tests = int(testsuite_elem.get('tests', '0'))
        failures = int(testsuite_elem.get('failures', '0'))
        errors = int(testsuite_elem.get('errors', '0'))
        skipped = int(testsuite_elem.get('skipped', '0'))
        time = float(testsuite_elem.get('time', '0'))
        
        test_cases = []
        for testcase_elem in testsuite_elem.findall('testcase'):
            test_cases.append(self._parse_testcase(testcase_elem))
        
        return TestSuite(name, tests, failures, errors, skipped, time, test_cases)
    
    def _parse_testcase(self, testcase_elem) -> TestCase:
        """Parse a single testcase element."""
        name = testcase_elem.get('name', 'Unknown')
        classname = testcase_elem.get('classname', 'Unknown')
        time = float(testcase_elem.get('time', '0'))
        
        # Determine test status
        failure_elem = testcase_elem.find('failure')
        error_elem = testcase_elem.find('error')
        skipped_elem = testcase_elem.find('skipped')
        
        if failure_elem is not None:
            status = 'failed'
            failure_message = failure_elem.get('message', 'No message')
            failure_details = failure_elem.text or 'No details'
            return TestCase(name, classname, time, status, failure_message, failure_details)
        elif error_elem is not None:
            status = 'error'
            failure_message = error_elem.get('message', 'No message')
            failure_details = error_elem.text or 'No details'
            return TestCase(name, classname, time, status, failure_message, failure_details)
        elif skipped_elem is not None:
            status = 'skipped'
            skipped_message = skipped_elem.get('message', 'Test skipped')
            return TestCase(name, classname, time, status, skipped_message=skipped_message)
        else:
            status = 'passed'
            return TestCase(name, classname, time, status)


class WordReportGenerator:
    """Generates Word document reports from JUnit test results."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles
        
        # Title style
        if 'Custom Title' not in styles:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(20)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        if 'Custom Heading' not in styles:
            heading_style = styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        
        # Summary box style
        if 'Summary Box' not in styles:
            summary_style = styles.add_style('Summary Box', WD_STYLE_TYPE.PARAGRAPH)
            summary_font = summary_style.font
            summary_font.name = 'Arial'
            summary_font.size = Pt(11)
            summary_style.paragraph_format.left_indent = Inches(0.5)
            summary_style.paragraph_format.right_indent = Inches(0.5)
    
    def generate_report(self, test_suites: List[TestSuite], output_file: Path):
        """Generate Word document report from test suites."""
        self._add_title()
        self._add_executive_summary(test_suites)
        
        for suite in test_suites:
            self._add_test_suite_section(suite)
        
        self._add_footer()
        self.doc.save(output_file)
    
    def _add_title(self):
        """Add document title."""
        title = self.doc.add_paragraph('Test Execution Report', 'Custom Title')
        
        # Add generation timestamp
        timestamp = self.doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = timestamp.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(10)
        run.font.italic = True
        
        self.doc.add_paragraph()  # Add spacing
    
    def _add_executive_summary(self, test_suites: List[TestSuite]):
        """Add executive summary section."""
        self.doc.add_paragraph('Executive Summary', 'Custom Heading')
        
        # Calculate totals
        total_tests = sum(suite.tests for suite in test_suites)
        total_failures = sum(suite.failures for suite in test_suites)
        total_errors = sum(suite.errors for suite in test_suites)
        total_skipped = sum(suite.skipped for suite in test_suites)
        total_passed = total_tests - total_failures - total_errors - total_skipped
        total_time = sum(suite.time for suite in test_suites)
        
        # Create summary table
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        # Set table data
        summary_data = [
            ('Total Test Suites', str(len(test_suites))),
            ('Total Tests', str(total_tests)),
            ('Passed', str(total_passed)),
            ('Failed', str(total_failures)),
            ('Errors', str(total_errors)),
            ('Skipped', str(total_skipped)),
            ('Total Execution Time', f'{total_time:.2f} seconds')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            
            # Style the cells
            for j in range(2):
                cell = table.cell(i, j)
                cell.paragraphs[0].runs[0].font.name = 'Arial'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if j == 0:  # Label column
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Add success rate
        success_rate = (total_passed / total_tests * 100) if total_tests > 0 else 0
        success_para = self.doc.add_paragraph()
        success_run = success_para.add_run(f'Success Rate: {success_rate:.1f}%')
        success_run.font.name = 'Arial'
        success_run.font.size = Pt(12)
        success_run.font.bold = True
        
        # Color code success rate
        if success_rate >= 95:
            success_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif success_rate >= 80:
            success_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            success_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        self.doc.add_paragraph()  # Add spacing
    
    def _add_test_suite_section(self, suite: TestSuite):
        """Add a section for a test suite."""
        self.doc.add_paragraph(f'Test Suite: {suite.name}', 'Custom Heading')
        
        # Suite summary
        suite_para = self.doc.add_paragraph()
        suite_para.add_run(f'Tests: {suite.tests} | ')
        
        passed_run = suite_para.add_run(f'Passed: {suite.tests - suite.failures - suite.errors - suite.skipped} | ')
        passed_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        if suite.failures > 0:
            failed_run = suite_para.add_run(f'Failed: {suite.failures} | ')
            failed_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        if suite.errors > 0:
            error_run = suite_para.add_run(f'Errors: {suite.errors} | ')
            error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        if suite.skipped > 0:
            skipped_run = suite_para.add_run(f'Skipped: {suite.skipped} | ')
            skipped_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        suite_para.add_run(f'Time: {suite.time:.2f}s')
        
        # Add failed/error test details
        failed_tests = [tc for tc in suite.test_cases if tc.status in ['failed', 'error']]
        if failed_tests:
            self.doc.add_paragraph('Failed/Error Tests:', style='List Bullet')
            for test in failed_tests:
                self._add_failed_test_details(test)
        
        # Add test case summary table if there are many tests
        if len(suite.test_cases) > 10:
            self._add_test_case_summary_table(suite.test_cases)
        else:
            self._add_detailed_test_cases(suite.test_cases)
        
        self.doc.add_paragraph()  # Add spacing
    
    def _add_failed_test_details(self, test_case: TestCase):
        """Add detailed information for failed tests."""
        # Test name
        test_para = self.doc.add_paragraph(style='List Bullet 2')
        test_run = test_para.add_run(f'{test_case.classname}.{test_case.name}')
        test_run.font.bold = True
        test_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Failure message
        if test_case.failure_message:
            msg_para = self.doc.add_paragraph(style='List Bullet 3')
            msg_para.add_run(f'Message: {test_case.failure_message}')
        
        # Failure details (truncated if too long)
        if test_case.failure_details:
            details = test_case.failure_details
            if len(details) > 500:
                details = details[:500] + '...'
            
            details_para = self.doc.add_paragraph(style='List Bullet 3')
            details_run = details_para.add_run(f'Details: {details}')
            details_run.font.name = 'Courier New'
            details_run.font.size = Pt(9)
    
    def _add_test_case_summary_table(self, test_cases: List[TestCase]):
        """Add a summary table for test cases (for large test suites)."""
        self.doc.add_paragraph('Test Cases Summary:')
        
        # Create table with headers
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Test Name', 'Class', 'Status', 'Time (s)']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add test case rows
        for test_case in test_cases:
            row_cells = table.add_row().cells
            row_cells[0].text = test_case.name
            row_cells[1].text = test_case.classname
            row_cells[2].text = test_case.status.upper()
            row_cells[3].text = f'{test_case.time:.3f}'
            
            # Color code status
            status_run = row_cells[2].paragraphs[0].runs[0]
            if test_case.status == 'passed':
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif test_case.status in ['failed', 'error']:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif test_case.status == 'skipped':
                status_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    
    def _add_detailed_test_cases(self, test_cases: List[TestCase]):
        """Add detailed list of test cases (for smaller test suites)."""
        if not test_cases:
            return
            
        self.doc.add_paragraph('Test Cases:')
        
        for test_case in test_cases:
            test_para = self.doc.add_paragraph(style='List Bullet')
            
            # Status icon
            if test_case.status == 'passed':
                test_para.add_run('✓ ')
                test_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif test_case.status in ['failed', 'error']:
                test_para.add_run('✗ ')
                test_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
            elif test_case.status == 'skipped':
                test_para.add_run('- ')
                test_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)  # Gray
            
            # Test name and details
            test_para.add_run(f'{test_case.classname}.{test_case.name} ({test_case.time:.3f}s)')
    
    def _add_footer(self):
        """Add document footer."""
        self.doc.add_page_break()
        footer_para = self.doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run('End of Test Report')
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True


def main():
    """Main function to run the script."""
    parser = argparse.ArgumentParser(
        description='Convert JUnit XML test results to Word document report',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python junit_to_word_report.py test-results.xml report.docx
    python junit_to_word_report.py junit.xml test_report.docx
        """
    )
    
    parser.add_argument('input_file', type=Path, help='Input JUnit XML file')
    parser.add_argument('output_file', type=Path, help='Output Word document file')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    # Validate input file
    if not args.input_file.exists():
        print(f"Error: Input file '{args.input_file}' does not exist.")
        sys.exit(1)
    
    if not args.input_file.suffix.lower() == '.xml':
        print(f"Warning: Input file '{args.input_file}' does not have .xml extension.")
    
    # Ensure output directory exists
    args.output_file.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        # Parse JUnit XML
        if args.verbose:
            print(f"Parsing JUnit XML file: {args.input_file}")
        
        parser = JUnitXMLParser()
        test_suites = parser.parse(args.input_file)
        
        if args.verbose:
            total_tests = sum(suite.tests for suite in test_suites)
            print(f"Found {len(test_suites)} test suite(s) with {total_tests} total tests")
        
        # Generate Word report
        if args.verbose:
            print(f"Generating Word report: {args.output_file}")
        
        generator = WordReportGenerator()
        generator.generate_report(test_suites, args.output_file)
        
        print(f"Report successfully generated: {args.output_file}")
        
    except Exception as e:
        print(f"Error: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()