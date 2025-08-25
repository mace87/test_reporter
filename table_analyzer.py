from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import qn
import os
import sys

class TableStyleAnalyzer:
    def __init__(self):
        self.alignment_names = {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER", 
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE"
        }
        
        self.table_alignment_names = {
            WD_TABLE_ALIGNMENT.LEFT: "LEFT",
            WD_TABLE_ALIGNMENT.CENTER: "CENTER",
            WD_TABLE_ALIGNMENT.RIGHT: "RIGHT"
        }

    def inches_to_points(self, inches_value):
        """Convert inches to points for easier reading"""
        if inches_value is None:
            return None
        return round(inches_value * 72, 1)  # 72 points per inch

    def rgb_to_hex(self, rgb_color):
        """Convert RGBColor to hex string"""
        if rgb_color is None:
            return None
        try:
            # Handle RGBColor object
            if hasattr(rgb_color, '__iter__') and len(rgb_color) == 3:
                # If it's already a tuple/list of RGB values
                r, g, b = rgb_color
                return f"#{r:02X}{g:02X}{b:02X}"
            elif isinstance(rgb_color, int):
                # If it's an integer representation
                return f"#{rgb_color:06X}"
            else:
                # If it's an RGBColor object, convert to int first
                rgb_int = int(rgb_color)
                return f"#{rgb_int:06X}"
        except (TypeError, ValueError, AttributeError):
            return str(rgb_color)

    def rgb_color_to_components(self, rgb_color):
        """Convert RGBColor to individual R, G, B components"""
        if rgb_color is None:
            return None, None, None
        
        try:
            # Convert RGBColor to integer
            if hasattr(rgb_color, '__iter__') and len(rgb_color) == 3:
                return rgb_color  # Already components
            
            # Convert RGBColor object to integer
            rgb_int = int(rgb_color)
            r = (rgb_int >> 16) & 0xFF
            g = (rgb_int >> 8) & 0xFF
            b = rgb_int & 0xFF
            return r, g, b
        except (TypeError, ValueError, AttributeError):
            return None, None, None

    def get_font_info(self, run):
        """Extract detailed font information from a run"""
        font_info = {}
        
        # Font name
        if run.font.name:
            font_info['name'] = run.font.name
        else:
            # Try to get the default font from the document
            font_info['name'] = "Default (typically Calibri)"
        
        # Font size
        if run.font.size:
            font_info['size_points'] = self.inches_to_points(run.font.size)
            font_info['size_inches'] = run.font.size
        else:
            font_info['size_points'] = "Default (typically 11pt)"
            font_info['size_inches'] = None
        
        # Font color
        if run.font.color.rgb:
            font_info['color'] = self.rgb_to_hex(run.font.color.rgb)
            font_info['color_components'] = self.rgb_color_to_components(run.font.color.rgb)
        else:
            font_info['color'] = "Default (typically black)"
            font_info['color_components'] = (None, None, None)
        
        # Font styling
        font_info['bold'] = run.bold if run.bold is not None else False
        font_info['italic'] = run.italic if run.italic is not None else False
        font_info['underline'] = run.underline if run.underline is not None else False
        
        # Additional font properties
        font_info['small_caps'] = run.font.small_caps if run.font.small_caps is not None else False
        font_info['all_caps'] = run.font.all_caps if run.font.all_caps is not None else False
        font_info['strike'] = run.font.strike if run.font.strike is not None else False
        font_info['subscript'] = run.font.subscript if run.font.subscript is not None else False
        font_info['superscript'] = run.font.superscript if run.font.superscript is not None else False
        
        return font_info

    def get_cell_background_color(self, cell):
        """Extract background color from cell"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        return f"#{fill}"
            return None
        except Exception:
            return None

    def get_cell_borders(self, cell):
        """Extract border information from cell"""
        borders = {}
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is not None:
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    for border_type in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(qn(f'w:{border_type}'))
                        if border is not None:
                            borders[border_type] = {
                                'style': border.get(qn('w:val'), 'none'),
                                'size': border.get(qn('w:sz'), '0'),
                                'color': border.get(qn('w:color'), 'auto')
                            }
            return borders if borders else None
        except Exception:
            return None

    def analyze_text_formatting(self, paragraph):
        """Analyze text formatting in a paragraph"""
        formatting = {
            'text': paragraph.text.strip(),
            'alignment': self.alignment_names.get(paragraph.alignment, str(paragraph.alignment)),
            'runs': []
        }
        
        for run in paragraph.runs:
            if run.text.strip():  # Only analyze runs with actual text
                font_info = self.get_font_info(run)
                run_format = {
                    'text': run.text,
                    'font': font_info
                }
                formatting['runs'].append(run_format)
        
        return formatting

    def check_merged_cells(self, table):
        """Detect merged cells in the table"""
        merged_info = []
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Check if this cell spans multiple columns or rows
                try:
                    tc = cell._tc
                    
                    # Check for column span (gridSpan)
                    tcPr = tc.tcPr
                    if tcPr is not None:
                        gridSpan = tcPr.find(qn('w:gridSpan'))
                        if gridSpan is not None:
                            span_value = int(gridSpan.get(qn('w:val'), '1'))
                            if span_value > 1:
                                merged_info.append({
                                    'row': row_idx,
                                    'col': col_idx,
                                    'col_span': span_value,
                                    'text': cell.text.strip()
                                })
                        
                        # Check for row span (vMerge)
                        vMerge = tcPr.find(qn('w:vMerge'))
                        if vMerge is not None:
                            val = vMerge.get(qn('w:val'))
                            if val == 'restart':
                                merged_info.append({
                                    'row': row_idx,
                                    'col': col_idx,
                                    'row_span_start': True,
                                    'text': cell.text.strip()
                                })
                            elif val is None:  # Continuation of merge
                                merged_info.append({
                                    'row': row_idx,
                                    'col': col_idx,
                                    'row_span_continue': True
                                })
                except Exception as e:
                    continue
        
        return merged_info

    def print_font_details(self, font_info, indent="      "):
        """Print detailed font information in a formatted way"""
        print(f"{indent}Font: {font_info['name']}")
        print(f"{indent}Font size: {font_info['size_points']} points")
        print(f"{indent}Font color: {font_info['color']}")
        
        # Style properties
        styles = []
        if font_info['bold']:
            styles.append("Bold")
        if font_info['italic']:
            styles.append("Italic")
        if font_info['underline']:
            styles.append("Underlined")
        if font_info['small_caps']:
            styles.append("Small Caps")
        if font_info['all_caps']:
            styles.append("All Caps")
        if font_info['strike']:
            styles.append("Strikethrough")
        if font_info['subscript']:
            styles.append("Subscript")
        if font_info['superscript']:
            styles.append("Superscript")
        
        if styles:
            print(f"{indent}Font styles: {', '.join(styles)}")
        else:
            print(f"{indent}Font styles: None")

    def analyze_table(self, table, table_index):
        """Analyze a single table and return detailed style information"""
        print(f"\n{'='*90}")
        print(f"TABLE {table_index + 1} ANALYSIS")
        print(f"{'='*90}")
        
        # Basic table information
        print(f"Dimensions: {len(table.rows)} rows × {len(table.columns)} columns")
        print(f"Table style: {table.style.name if table.style else 'None'}")
        print(f"Table alignment: {self.table_alignment_names.get(table.alignment, str(table.alignment))}")
        
        # Check for merged cells
        merged_cells = self.check_merged_cells(table)
        if merged_cells:
            print(f"\nMERGED CELLS DETECTED:")
            for merge in merged_cells:
                if 'col_span' in merge:
                    print(f"  Row {merge['row']}, Col {merge['col']}: spans {merge['col_span']} columns - '{merge['text']}'")
                elif 'row_span_start' in merge:
                    print(f"  Row {merge['row']}, Col {merge['col']}: row merge start - '{merge['text']}'")
                elif 'row_span_continue' in merge:
                    print(f"  Row {merge['row']}, Col {merge['col']}: row merge continuation")
        
        # Analyze each cell
        print(f"\nCELL-BY-CELL ANALYSIS:")
        print("-" * 90)
        
        for row_idx, row in enumerate(table.rows):
            print(f"\nROW {row_idx}:")
            
            # Check row height
            try:
                row_height = self.inches_to_points(row.height)
                if row_height:
                    print(f"  Row height: {row_height} points")
            except:
                pass
            
            for col_idx, cell in enumerate(row.cells):
                print(f"\n  CELL [{row_idx}, {col_idx}]:")
                
                # Cell dimensions
                try:
                    width = self.inches_to_points(cell.width)
                    if width:
                        print(f"    Width: {width} points")
                except:
                    pass
                
                # Background color
                bg_color = self.get_cell_background_color(cell)
                if bg_color:
                    print(f"    Background color: {bg_color}")
                
                # Borders
                borders = self.get_cell_borders(cell)
                if borders:
                    print(f"    Borders:")
                    for border_type, border_info in borders.items():
                        color = f"#{border_info['color']}" if border_info['color'] != 'auto' else border_info['color']
                        print(f"      {border_type}: {border_info['style']}, size: {border_info['size']}, color: {color}")
                
                # Text formatting for each paragraph
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():  # Only analyze paragraphs with text
                        formatting = self.analyze_text_formatting(paragraph)
                        print(f"    Paragraph {para_idx}:")
                        print(f"      Text: '{formatting['text']}'")
                        print(f"      Alignment: {formatting['alignment']}")
                        
                        for run_idx, run_format in enumerate(formatting['runs']):
                            print(f"      Run {run_idx}:")
                            print(f"        Text: '{run_format['text']}'")
                            self.print_font_details(run_format['font'], "        ")

    def generate_replication_code(self, table, table_index):
        """Generate Python code snippet to replicate the table style"""
        print(f"\n{'='*90}")
        print(f"PYTHON CODE TO REPLICATE TABLE {table_index + 1}")
        print(f"{'='*90}")
        
        rows = len(table.rows)
        cols = len(table.columns)
        
        code = f"""
# Import required modules
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# Create table with {rows} rows and {cols} columns
table = doc.add_table(rows={rows}, cols={cols})

# Set table style and alignment
table.style = '{table.style.name if table.style else 'Table Grid'}'
table.alignment = WD_TABLE_ALIGNMENT.{self.table_alignment_names.get(table.alignment, 'LEFT')}

# Style individual cells
"""
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():  # Only generate code for cells with content
                    code += f"\n# Cell [{row_idx}, {col_idx}]\n"
                    code += f"cell = table.rows[{row_idx}].cells[{col_idx}]\n"
                    code += f"cell.text = '{cell.text.strip()}'\n"
                    
                    # Background color
                    bg_color = self.get_cell_background_color(cell)
                    if bg_color:
                        hex_color = bg_color.replace('#', '').upper()
                        code += f"""
# Set background color
cell_xml = cell._tc
cell_properties = cell_xml.get_or_add_tcPr()
shading = OxmlElement('w:shd')
shading.set(qn('w:val'), 'clear')
shading.set(qn('w:color'), 'auto')
shading.set(qn('w:fill'), '{hex_color}')
cell_properties.append(shading)
"""
                    
                    # Text formatting
                    if cell.paragraphs:
                        paragraph = cell.paragraphs[0]
                        formatting_code = []
                        
                        if paragraph.alignment is not None:
                            alignment_name = self.alignment_names.get(paragraph.alignment, 'LEFT')
                            formatting_code.append(f"paragraph = cell.paragraphs[0]")
                            formatting_code.append(f"paragraph.alignment = WD_ALIGN_PARAGRAPH.{alignment_name}")
                        
                        if paragraph.runs:
                            run = paragraph.runs[0]
                            font_info = self.get_font_info(run)
                            
                            run_code = ["run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()"]
                            
                            # Font properties
                            if font_info['name'] and font_info['name'] != "Default (typically Calibri)":
                                run_code.append(f"run.font.name = '{font_info['name']}'")
                            
                            if font_info['size_inches']:
                                run_code.append(f"run.font.size = Inches({font_info['size_inches']})")
                            
                            # Handle font color with proper RGB component extraction
                            if (font_info['color'] and 
                                font_info['color'] != "Default (typically black)" and 
                                font_info['color_components'][0] is not None):
                                r, g, b = font_info['color_components']
                                run_code.append(f"run.font.color.rgb = RGBColor({r}, {g}, {b})")
                            
                            # Style properties
                            if font_info['bold']:
                                run_code.append("run.font.bold = True")
                            if font_info['italic']:
                                run_code.append("run.font.italic = True")
                            if font_info['underline']:
                                run_code.append("run.font.underline = True")
                            if font_info['small_caps']:
                                run_code.append("run.font.small_caps = True")
                            if font_info['all_caps']:
                                run_code.append("run.font.all_caps = True")
                            if font_info['strike']:
                                run_code.append("run.font.strike = True")
                            if font_info['subscript']:
                                run_code.append("run.font.subscript = True")
                            if font_info['superscript']:
                                run_code.append("run.font.superscript = True")
                            
                            if len(run_code) > 1:  # More than just the run creation line
                                formatting_code.extend(run_code)
                        
                        if formatting_code:
                            code += "\n" + "\n".join(formatting_code) + "\n"
        
        # Add merged cells code
        merged_cells = self.check_merged_cells(table)
        if merged_cells:
            code += "\n# Merge cells\n"
            for merge in merged_cells:
                if 'col_span' in merge:
                    row, col, span = merge['row'], merge['col'], merge['col_span']
                    end_col = col + span - 1
                    code += f"table.rows[{row}].cells[{col}].merge(table.rows[{row}].cells[{end_col}])\n"
        
        print(code)

    def analyze_document(self, file_path):
        """Analyze all tables in a Word document"""
        if not os.path.exists(file_path):
            print(f"Error: File '{file_path}' not found!")
            return
        
        try:
            print(f"Analyzing document: {file_path}")
            doc = Document(file_path)
            
            if not doc.tables:
                print("No tables found in the document!")
                return
            
            print(f"Found {len(doc.tables)} table(s) in the document.")
            
            for i, table in enumerate(doc.tables):
                self.analyze_table(table, i)
                self.generate_replication_code(table, i)
                
                if i < len(doc.tables) - 1:
                    input("\nPress Enter to continue to the next table...")
            
            print(f"\n{'='*90}")
            print("ANALYSIS COMPLETE")
            print(f"{'='*90}")
            
        except Exception as e:
            print(f"Error analyzing document: {str(e)}")
            import traceback
            traceback.print_exc()

def main():
    """Main function to run the table analyzer"""
    analyzer = TableStyleAnalyzer()
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = input("Enter the path to the Word document (.docx): ").strip().strip('"')
    
    if not file_path:
        print("No file path provided!")
        return
    
    analyzer.analyze_document(file_path)

if __name__ == "__main__":
    main()