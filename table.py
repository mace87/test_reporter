from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def create_styled_table():
    # Create a new document
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Sales Report - Q3 2024', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create a table with 4 rows and 4 columns
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table width
    table.autofit = False
    table.allow_autofit = False
    
    # Define column widths
    for i, width in enumerate([Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.2)]):
        for row in table.rows:
            row.cells[i].width = width
    
    # Merge cells in the header row for a spanning title
    header_cell = table.cell(0, 0)
    header_cell_end = table.cell(0, 3)
    header_cell.merge(header_cell_end)
    
    # Set the merged header cell content
    header_paragraph = header_cell.paragraphs[0]
    header_run = header_paragraph.runs[0] if header_paragraph.runs else header_paragraph.add_run()
    header_run.text = "Quarterly Sales Performance Summary"
    header_run.font.bold = True
    header_run.font.size = Pt(14)
    header_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Color the header row background
    def set_cell_background_color(cell, color):
        """Set background color for a cell"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            print(f"Warning: Could not set background color: {e}")
    
    set_cell_background_color(header_cell, "4472C4")  # Blue background
    
    # Add column headers in row 1
    headers = ["Product Category", "Units Sold", "Revenue ($)", "Growth (%)"]
    header_row = table.rows[1]
    
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background_color(cell, "5B9BD5")  # Lighter blue background
    
    # Sample data
    data = [
        ["Electronics", "12,450", "$1,245,000", "+15.2%"],
        ["Clothing", "8,230", "$823,000", "+8.7%"]
    ]
    
    # Fill data rows
    for row_idx, row_data in enumerate(data, start=2):  # Start from row 2
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(cell_data)
            
            # Style first column (Product Category) - bold
            if col_idx == 0:
                run.font.bold = True
                run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Style second and third columns (numbers) - monospace font
            elif col_idx == 1 or col_idx == 2:
                run.font.name = 'Courier New'  # Monospace font
                run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Style fourth column (percentage) - centered
            else:
                run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add color coding for growth percentage
                if "+" in cell_data:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green for positive
                elif "-" in cell_data:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red for negative
    
    # Add some spacing after the table
    doc.add_paragraph()
    
    # Add a note
    note = doc.add_paragraph("Note: All figures are preliminary and subject to final audit.")
    note_run = note.runs[0]
    note_run.font.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray text
    
    # Save the document
    doc.save('styled_table_example.docx')
    print("Document saved as 'styled_table_example.docx'")

if __name__ == "__main__":
    create_styled_table()